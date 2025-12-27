"""
Generate Word document templates for personal productivity system.
Creates 6 templates that can be imported into OneNote or used directly.
"""

from docx import Document
from pathlib import Path


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading


def add_line(doc):
    doc.add_paragraph("_" * 60)


def add_checkbox_item(doc, text=""):
    p = doc.add_paragraph()
    p.add_run("☐ " + text)


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def create_monthly_template(output_dir):
    doc = Document()
    add_heading(doc, "MONTHLY REVIEW")

    doc.add_paragraph("Month: _______________________")
    doc.add_paragraph()

    add_heading(doc, "Goals This Month", level=2)
    for _ in range(3):
        add_checkbox_item(doc)
    doc.add_paragraph()

    add_heading(doc, "Sessions Needed to Hit Goals", level=2)
    for _ in range(5):
        add_bullet(doc, "")
    doc.add_paragraph()

    add_heading(doc, "Last Month Reflection", level=2)
    doc.add_paragraph("What worked:")
    doc.add_paragraph()
    doc.add_paragraph("What didn't:")
    doc.add_paragraph()

    add_heading(doc, "Theme/Focus", level=2)
    doc.add_paragraph()

    doc.save(output_dir / "01_monthly_review.docx")
    print("Created: 01_monthly_review.docx")


def create_session_backlog_template(output_dir):
    doc = Document()
    add_heading(doc, "SESSION BACKLOG")
    doc.add_paragraph("Planned work blocks - schedule these into your week")
    doc.add_paragraph()

    add_heading(doc, "Priority Sessions", level=2)

    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"

    headers = ["Priority", "Session Name", "Est. Time", "Status"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    for row in table.rows[1:]:
        row.cells[0].text = ""
        row.cells[3].text = "Not scheduled"

    doc.add_paragraph()
    add_heading(doc, "Completed Sessions", level=2)
    doc.add_paragraph("(Move here when done)")
    for _ in range(3):
        add_bullet(doc, "")

    doc.save(output_dir / "02_session_backlog.docx")
    print("Created: 02_session_backlog.docx")


def create_weekly_template(output_dir):
    doc = Document()
    add_heading(doc, "WEEKLY PLAN")

    doc.add_paragraph("Week of: _______________________")
    doc.add_paragraph()

    add_heading(doc, "Goal for This Week", level=2)
    doc.add_paragraph()

    add_heading(doc, "Sessions Scheduled", level=2)

    days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Weekend"]
    for day in days:
        doc.add_paragraph(f"{day}:")

    doc.add_paragraph()
    add_heading(doc, "Last Week Review", level=2)
    doc.add_paragraph("What got done:")
    doc.add_paragraph()
    doc.add_paragraph("What got blocked:")
    doc.add_paragraph()
    doc.add_paragraph("Carry forward:")
    doc.add_paragraph()

    doc.save(output_dir / "03_weekly_plan.docx")
    print("Created: 03_weekly_plan.docx")


def create_daily_template(output_dir):
    doc = Document()
    add_heading(doc, "DAILY PLAN")

    doc.add_paragraph("Date: _______________________")
    doc.add_paragraph()

    add_heading(doc, "Today's Sessions", level=2)
    doc.add_paragraph("1.")
    doc.add_paragraph("2.")
    doc.add_paragraph("3.")
    doc.add_paragraph()

    add_heading(doc, "Must Win Today", level=2)
    doc.add_paragraph()

    add_heading(doc, "Notes/Adjustments", level=2)
    doc.add_paragraph()

    add_heading(doc, "End of Day - Did I Win?", level=2)
    p = doc.add_paragraph()
    p.add_run("☐ Yes     ☐ No")
    doc.add_paragraph()
    doc.add_paragraph("If no, why:")

    doc.save(output_dir / "04_daily_plan.docx")
    print("Created: 04_daily_plan.docx")


def create_capture_log_template(output_dir):
    doc = Document()
    add_heading(doc, "CAPTURE LOG")
    doc.add_paragraph("One-liners while working - process later")
    doc.add_paragraph()

    table = doc.add_table(rows=11, cols=3)
    table.style = "Table Grid"

    headers = ["Date", "What I Noticed", "Type"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    doc.add_paragraph()
    doc.add_paragraph("Types: bug / idea / friction / win / question")

    doc.save(output_dir / "05_capture_log.docx")
    print("Created: 05_capture_log.docx")


def create_content_bank_template(output_dir):
    doc = Document()
    add_heading(doc, "CONTENT IDEA BANK")
    doc.add_paragraph("Feeds LinkedIn/Medium campaign")
    doc.add_paragraph()

    add_heading(doc, "Ideas", level=2)
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"

    headers = ["Date", "Idea", "Context", "Format"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    doc.add_paragraph()
    doc.add_paragraph("Formats: Post / Article / Thread")
    doc.add_paragraph()

    add_heading(doc, "Ready to Write", level=2)
    for _ in range(3):
        add_bullet(doc, "")

    add_heading(doc, "Drafted", level=2)
    for _ in range(3):
        add_bullet(doc, "")

    add_heading(doc, "Published", level=2)
    for _ in range(3):
        add_bullet(doc, "")

    doc.save(output_dir / "06_content_idea_bank.docx")
    print("Created: 06_content_idea_bank.docx")


def main():
    output_dir = Path(__file__).parent.parent / "templates"
    output_dir.mkdir(exist_ok=True)

    print(f"Creating templates in: {output_dir}")
    print()

    create_monthly_template(output_dir)
    create_session_backlog_template(output_dir)
    create_weekly_template(output_dir)
    create_daily_template(output_dir)
    create_capture_log_template(output_dir)
    create_content_bank_template(output_dir)

    print()
    print(f"Done! 6 templates created in {output_dir}")
    print("You can now import these into OneNote or use them directly.")


if __name__ == "__main__":
    main()
