"""
Generate the complete Data Engineering Planning System.
Creates Word docs for narrative planning and Excel workbooks for structured data.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from pathlib import Path


# Styling constants
HEADER_FILL = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
SUBHEADER_FILL = PatternFill(start_color="D6DCE5", end_color="D6DCE5", fill_type="solid")
THIN_BORDER = Border(
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin"),
)


def style_header_row(ws, row_num, num_cols):
    """Apply header styling to a row."""
    for col in range(1, num_cols + 1):
        cell = ws.cell(row=row_num, column=col)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = THIN_BORDER


def style_data_cell(cell, wrap=True):
    """Apply standard styling to a data cell."""
    cell.border = THIN_BORDER
    cell.alignment = Alignment(vertical="top", wrap_text=wrap)


def set_column_widths(ws, widths):
    """Set column widths from a list."""
    for i, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = width


def add_section_header(doc, text):
    """Add a styled section header to a Word doc."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def add_thinking_prompt(doc, prompt):
    """Add a thinking prompt - the key to forcing planning."""
    p = doc.add_paragraph()
    run = p.add_run("→ " + prompt)
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    return p


# =============================================================================
# WORD TEMPLATES
# =============================================================================


def create_monthly_goals(output_dir):
    """Monthly Goals - High-level direction setting."""
    doc = Document()

    doc.add_heading("MONTHLY GOALS", 0)
    doc.add_paragraph("Month: _______________________")
    doc.add_paragraph()

    add_section_header(doc, "REFLECT ON LAST MONTH")
    add_thinking_prompt(doc, "What did I ship? What am I proud of?")
    doc.add_paragraph()
    doc.add_paragraph()
    add_thinking_prompt(doc, "What didn't get done? Why?")
    doc.add_paragraph()
    doc.add_paragraph()
    add_thinking_prompt(doc, "What lesson am I carrying forward?")
    doc.add_paragraph()
    doc.add_paragraph()

    add_section_header(doc, "THIS MONTH'S FOCUS")
    add_thinking_prompt(doc, "What is the ONE big thing that matters most this month?")
    doc.add_paragraph()
    doc.add_paragraph()

    add_section_header(doc, "GOALS BY AREA")

    doc.add_paragraph("Work (Ingredion):", style="List Bullet")
    add_thinking_prompt(doc, "What pipelines am I migrating/building? What must ship?")
    doc.add_paragraph()

    doc.add_paragraph("Odibi Development:", style="List Bullet")
    add_thinking_prompt(doc, "What features/fixes? What version am I targeting?")
    doc.add_paragraph()

    doc.add_paragraph("Marketing Campaign:", style="List Bullet")
    add_thinking_prompt(doc, "How many posts/articles? What themes?")
    doc.add_paragraph()

    add_section_header(doc, "SESSIONS NEEDED")
    add_thinking_prompt(doc, "What specific work blocks do I need to schedule to hit these goals?")
    table = doc.add_table(rows=8, cols=4)
    table.style = "Table Grid"
    headers = ["Session Name", "Type", "Est. Hours", "Priority"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    doc.add_paragraph()
    add_section_header(doc, "RISKS / BLOCKERS")
    add_thinking_prompt(doc, "What could stop me? What do I need from others?")
    doc.add_paragraph()

    doc.save(output_dir / "01_monthly_goals.docx")
    print("Created: 01_monthly_goals.docx")


def create_weekly_plan(output_dir):
    """Weekly Plan - Allocate sessions to days."""
    doc = Document()

    doc.add_heading("WEEKLY PLAN", 0)
    doc.add_paragraph("Week of: _______________________")
    doc.add_paragraph()

    add_section_header(doc, "LAST WEEK REVIEW")
    add_thinking_prompt(doc, "What got done?")
    doc.add_paragraph()
    add_thinking_prompt(doc, "What got blocked? Why?")
    doc.add_paragraph()
    add_thinking_prompt(doc, "What's carrying forward?")
    doc.add_paragraph()

    add_section_header(doc, "THIS WEEK'S WIN")
    add_thinking_prompt(doc, "If I only accomplish ONE thing this week, what must it be?")
    doc.add_paragraph()
    doc.add_paragraph()

    add_section_header(doc, "SESSIONS THIS WEEK")
    add_thinking_prompt(doc, "Pull from your session backlog. Be realistic about time.")

    table = doc.add_table(rows=7, cols=4)
    table.style = "Table Grid"
    headers = ["Day", "Session(s)", "Time Block", "Must Win"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Weekend"]
    for i, day in enumerate(days):
        table.rows[i + 1].cells[0].text = day

    doc.add_paragraph()
    add_section_header(doc, "BLOCKERS TO CLEAR")
    add_thinking_prompt(doc, "What must I resolve early in the week to unblock later work?")
    doc.add_paragraph()

    add_section_header(doc, "NOTES")
    doc.add_paragraph()

    doc.save(output_dir / "03_weekly_plan.docx")
    print("Created: 03_weekly_plan.docx")


def create_daily_plan(output_dir):
    """Daily Plan - Focus for today."""
    doc = Document()

    doc.add_heading("DAILY PLAN", 0)
    doc.add_paragraph("Date: _______________________")
    doc.add_paragraph()

    add_section_header(doc, "TODAY'S WIN")
    add_thinking_prompt(doc, "What ONE thing, if completed, makes today a success?")
    doc.add_paragraph()
    doc.add_paragraph()

    add_section_header(doc, "SESSIONS TODAY")
    add_thinking_prompt(doc, "Which sessions from your weekly plan are you tackling?")

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Session", "Time Block", "Definition of Done"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    doc.add_paragraph()
    add_section_header(doc, "BEFORE I START")
    doc.add_paragraph("☐ Do I have everything I need? (data access, credentials, docs)")
    doc.add_paragraph("☐ Is my environment ready? (Odibi, connections, test data)")
    doc.add_paragraph("☐ Have I removed distractions? (notifications, tabs, phone)")

    doc.add_paragraph()
    add_section_header(doc, "END OF DAY")
    add_thinking_prompt(doc, "Did I win today?")
    doc.add_paragraph("☐ Yes")
    doc.add_paragraph("☐ No - Why:")
    doc.add_paragraph()
    add_thinking_prompt(doc, "What's the first thing I'm doing tomorrow?")
    doc.add_paragraph()

    doc.save(output_dir / "04_daily_plan.docx")
    print("Created: 04_daily_plan.docx")


# =============================================================================
# EXCEL TEMPLATES
# =============================================================================


def create_session_backlog(output_dir):
    """Session Backlog - All planned work blocks."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Session Backlog"

    # Headers
    headers = [
        "Session ID",
        "Session Name",
        "Type",
        "Layer",
        "Pipeline/Node",
        "Est. Hours",
        "Priority",
        "Status",
        "Scheduled Week",
        "Notes",
    ]

    for col, header in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=header)

    style_header_row(ws, 1, len(headers))

    # Example data types in row 2 as guidance
    examples = [
        "S-001",
        "Migrate sales orders",
        "Migration",
        "Bronze",
        "brz_sales_orders",
        "4",
        "High",
        "Not Started",
        "Week 2",
        "Depends on SAP access",
    ]
    for col, ex in enumerate(examples, 1):
        cell = ws.cell(row=2, column=col, value=ex)
        cell.font = Font(italic=True, color="808080")
        style_data_cell(cell)

    # Add more empty rows
    for row in range(3, 52):
        for col in range(1, len(headers) + 1):
            style_data_cell(ws.cell(row=row, column=col))

    set_column_widths(ws, [10, 30, 12, 10, 25, 10, 10, 12, 15, 30])

    # Add dropdown guidance sheet
    ws2 = wb.create_sheet("Dropdowns")
    ws2["A1"] = "Type Options"
    ws2["A2"] = "Migration"
    ws2["A3"] = "New Build"
    ws2["A4"] = "Bug Fix"
    ws2["A5"] = "Optimization"
    ws2["A6"] = "Documentation"
    ws2["A7"] = "Testing"

    ws2["B1"] = "Layer Options"
    ws2["B2"] = "Bronze"
    ws2["B3"] = "Silver"
    ws2["B4"] = "Gold"
    ws2["B5"] = "Semantic"

    ws2["C1"] = "Status Options"
    ws2["C2"] = "Not Started"
    ws2["C3"] = "In Progress"
    ws2["C4"] = "Blocked"
    ws2["C5"] = "Done"

    ws2["D1"] = "Priority Options"
    ws2["D2"] = "High"
    ws2["D3"] = "Medium"
    ws2["D4"] = "Low"

    wb.save(output_dir / "02_session_backlog.xlsx")
    print("Created: 02_session_backlog.xlsx")


def create_master_inventory(output_dir):
    """Master Inventory - All nodes across all pipelines."""
    wb = Workbook()

    # ---------- OVERVIEW SHEET ----------
    ws_overview = wb.active
    ws_overview.title = "Overview"

    ws_overview["A1"] = "MASTER PIPELINE INVENTORY"
    ws_overview["A1"].font = Font(bold=True, size=16)

    ws_overview["A3"] = "Last Updated:"
    ws_overview["B3"] = "_____________"

    ws_overview["A5"] = "Summary"
    ws_overview["A5"].font = Font(bold=True, size=12)

    summary_data = [
        ["Layer", "Total Nodes", "Migrated", "In Progress", "Not Started"],
        ["Bronze", "", "", "", ""],
        ["Silver", "", "", "", ""],
        ["Gold", "", "", "", ""],
        ["TOTAL", "", "", "", ""],
    ]

    for row_idx, row_data in enumerate(summary_data, 6):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws_overview.cell(row=row_idx, column=col_idx, value=value)
            if row_idx == 6:
                cell.fill = HEADER_FILL
                cell.font = HEADER_FONT
            cell.border = THIN_BORDER

    # ---------- BRONZE SHEET ----------
    ws_bronze = wb.create_sheet("Bronze Nodes")
    bronze_headers = [
        "Node ID",
        "Node Name",
        "Source System",
        "Source Table(s)",
        "Target Table",
        "Pattern",
        "Key Columns",
        "Row Count (Est)",
        "Schedule",
        "Status",
        "Priority",
        "Dependencies",
        "Notes",
    ]

    for col, header in enumerate(bronze_headers, 1):
        ws_bronze.cell(row=1, column=col, value=header)
    style_header_row(ws_bronze, 1, len(bronze_headers))

    # Example row
    bronze_example = [
        "BRZ-001",
        "brz_sales_orders",
        "SAP",
        "VBAK, VBAP",
        "bronze.sales_orders",
        "Append",
        "order_id, line_id",
        "2,000,000",
        "Daily 6am",
        "Migrated",
        "High",
        "None",
        "Legacy SSIS",
    ]
    for col, val in enumerate(bronze_example, 1):
        cell = ws_bronze.cell(row=2, column=col, value=val)
        cell.font = Font(italic=True, color="808080")
        style_data_cell(cell)

    for row in range(3, 52):
        for col in range(1, len(bronze_headers) + 1):
            style_data_cell(ws_bronze.cell(row=row, column=col))

    set_column_widths(ws_bronze, [10, 25, 15, 25, 25, 12, 25, 15, 12, 12, 10, 25, 30])

    # ---------- SILVER SHEET ----------
    ws_silver = wb.create_sheet("Silver Nodes")
    silver_headers = [
        "Node ID",
        "Node Name",
        "Source Node(s)",
        "Target Table",
        "Pattern",
        "Transform Type",
        "Key Columns",
        "Business Logic",
        "Row Count (Est)",
        "Schedule",
        "Status",
        "Priority",
        "Notes",
    ]

    for col, header in enumerate(silver_headers, 1):
        ws_silver.cell(row=1, column=col, value=header)
    style_header_row(ws_silver, 1, len(silver_headers))

    silver_example = [
        "SLV-001",
        "slv_sales_enriched",
        "brz_sales_orders, brz_customers",
        "silver.sales_enriched",
        "Merge",
        "Join + Enrich",
        "order_id",
        "Join on customer_id, add customer name",
        "1,800,000",
        "Daily 7am",
        "In Progress",
        "High",
        "Replaces report_sales_v2",
    ]
    for col, val in enumerate(silver_example, 1):
        cell = ws_silver.cell(row=2, column=col, value=val)
        cell.font = Font(italic=True, color="808080")
        style_data_cell(cell)

    for row in range(3, 52):
        for col in range(1, len(silver_headers) + 1):
            style_data_cell(ws_silver.cell(row=row, column=col))

    set_column_widths(ws_silver, [10, 25, 30, 25, 12, 15, 20, 35, 15, 12, 12, 10, 30])

    # ---------- GOLD SHEET ----------
    ws_gold = wb.create_sheet("Gold Nodes")
    gold_headers = [
        "Node ID",
        "Node Name",
        "Source Node(s)",
        "Target Table",
        "Pattern",
        "Model Type",
        "Grain",
        "Key Measures/Attributes",
        "Consumers",
        "Schedule",
        "Status",
        "Priority",
        "Notes",
    ]

    for col, header in enumerate(gold_headers, 1):
        ws_gold.cell(row=1, column=col, value=header)
    style_header_row(ws_gold, 1, len(gold_headers))

    gold_example = [
        "GLD-001",
        "dim_customer",
        "slv_customers",
        "gold.dim_customer",
        "Dimension",
        "SCD2",
        "customer_id",
        "name, region, segment, valid_from, valid_to",
        "Sales Dashboard, Finance Report",
        "Daily 8am",
        "Not Started",
        "Medium",
        "",
    ]
    for col, val in enumerate(gold_example, 1):
        cell = ws_gold.cell(row=2, column=col, value=val)
        cell.font = Font(italic=True, color="808080")
        style_data_cell(cell)

    for row in range(3, 32):
        for col in range(1, len(gold_headers) + 1):
            style_data_cell(ws_gold.cell(row=row, column=col))

    set_column_widths(ws_gold, [10, 25, 30, 25, 12, 12, 15, 35, 30, 12, 12, 10, 30])

    # ---------- REFERENCE SHEET ----------
    ws_ref = wb.create_sheet("Reference")
    ws_ref["A1"] = "Status Values"
    ws_ref["A1"].font = Font(bold=True)
    ws_ref["A2"] = "Not Started"
    ws_ref["A3"] = "In Progress"
    ws_ref["A4"] = "Migrated"
    ws_ref["A5"] = "Blocked"
    ws_ref["A6"] = "Deprecated"

    ws_ref["C1"] = "Pattern Values"
    ws_ref["C1"].font = Font(bold=True)
    ws_ref["C2"] = "Append"
    ws_ref["C3"] = "Merge"
    ws_ref["C4"] = "SCD2"
    ws_ref["C5"] = "Dimension"
    ws_ref["C6"] = "Fact"
    ws_ref["C7"] = "Aggregation"
    ws_ref["C8"] = "DateDimension"

    ws_ref["E1"] = "Priority Values"
    ws_ref["E1"].font = Font(bold=True)
    ws_ref["E2"] = "High"
    ws_ref["E3"] = "Medium"
    ws_ref["E4"] = "Low"

    wb.save(output_dir / "05_master_inventory.xlsx")
    print("Created: 05_master_inventory.xlsx")


def create_project_planning(output_dir):
    """Project Planning Workbook - Detailed planning for migrations/new builds."""
    wb = Workbook()

    # ---------- OVERVIEW SHEET ----------
    ws_overview = wb.active
    ws_overview.title = "Overview"

    ws_overview["A1"] = "PROJECT PLANNING WORKBOOK"
    ws_overview["A1"].font = Font(bold=True, size=16)

    questions = [
        ("A3", "Project Name:"),
        ("A4", "Target Node(s):"),
        ("A5", "Layer:"),
        ("A6", "Pattern:"),
        ("A8", "WHAT AM I BUILDING/MIGRATING?"),
        ("A9", "→ Describe in plain English what this pipeline does:"),
        ("A12", "WHY DOES THIS EXIST?"),
        ("A13", "→ What business question does it answer? Who uses it?"),
        ("A16", "WHAT DOES SUCCESS LOOK LIKE?"),
        ("A17", "→ How will I know it's working correctly?"),
        ("A20", "OLD APPROACH (if migration):"),
        ("A21", "→ Paste or describe the old SQL/logic here:"),
        ("A30", "ODIBI APPROACH:"),
        ("A31", "→ Which pattern? What config options?"),
    ]

    for cell, text in questions:
        ws_overview[cell] = text
        if text.startswith("→"):
            ws_overview[cell].font = Font(italic=True, color="666666")
        elif text.isupper() or text.endswith(":") and not text.startswith("→"):
            ws_overview[cell].font = Font(bold=True)

    ws_overview.column_dimensions["A"].width = 100

    # ---------- SOURCES SHEET ----------
    ws_sources = wb.create_sheet("Sources")

    ws_sources["A1"] = "SOURCE INVENTORY"
    ws_sources["A1"].font = Font(bold=True, size=14)
    ws_sources["A2"] = "→ List ALL sources this pipeline reads from"
    ws_sources["A2"].font = Font(italic=True, color="666666")

    source_headers = [
        "Source #",
        "Source Name",
        "Type",
        "Location/Connection",
        "Key Columns",
        "Columns Needed",
        "Filter Conditions",
        "Row Count (Est)",
        "Update Frequency",
        "Notes",
    ]

    for col, header in enumerate(source_headers, 1):
        ws_sources.cell(row=4, column=col, value=header)
    style_header_row(ws_sources, 4, len(source_headers))

    for row in range(5, 20):
        ws_sources.cell(row=row, column=1, value=row - 4)
        for col in range(1, len(source_headers) + 1):
            style_data_cell(ws_sources.cell(row=row, column=col))

    set_column_widths(ws_sources, [10, 25, 12, 30, 25, 35, 30, 15, 15, 30])

    # ---------- COLUMN MAPPINGS SHEET ----------
    ws_mappings = wb.create_sheet("Column Mappings")

    ws_mappings["A1"] = "COLUMN MAPPINGS"
    ws_mappings["A1"].font = Font(bold=True, size=14)
    ws_mappings["A2"] = (
        "→ For each OUTPUT column, where does it come from and how is it transformed?"
    )
    ws_mappings["A2"].font = Font(italic=True, color="666666")

    mapping_headers = [
        "Target Column",
        "Data Type",
        "Source Table",
        "Source Column",
        "Transformation",
        "Nullable?",
        "Default Value",
        "Validation Rule",
        "Notes",
    ]

    for col, header in enumerate(mapping_headers, 1):
        ws_mappings.cell(row=4, column=col, value=header)
    style_header_row(ws_mappings, 4, len(mapping_headers))

    # Example row
    mapping_example = [
        "customer_name",
        "STRING",
        "brz_customers",
        "CUST_NAME",
        "TRIM(UPPER(source))",
        "No",
        "N/A",
        "NOT NULL, len > 0",
        "Standardize case",
    ]
    for col, val in enumerate(mapping_example, 1):
        cell = ws_mappings.cell(row=5, column=col, value=val)
        cell.font = Font(italic=True, color="808080")
        style_data_cell(cell)

    for row in range(6, 50):
        for col in range(1, len(mapping_headers) + 1):
            style_data_cell(ws_mappings.cell(row=row, column=col))

    set_column_widths(ws_mappings, [20, 12, 20, 20, 35, 10, 15, 25, 30])

    # ---------- JOINS SHEET ----------
    ws_joins = wb.create_sheet("Joins")

    ws_joins["A1"] = "JOIN LOGIC"
    ws_joins["A1"].font = Font(bold=True, size=14)
    ws_joins["A2"] = "→ Document every join. Be explicit about keys and join types."
    ws_joins["A2"].font = Font(italic=True, color="666666")

    join_headers = [
        "Join #",
        "Left Table",
        "Right Table",
        "Join Type",
        "Left Key(s)",
        "Right Key(s)",
        "Expected Cardinality",
        "Notes",
    ]

    for col, header in enumerate(join_headers, 1):
        ws_joins.cell(row=4, column=col, value=header)
    style_header_row(ws_joins, 4, len(join_headers))

    join_example = [
        "1",
        "brz_sales_orders",
        "brz_customers",
        "LEFT",
        "customer_id",
        "customer_id",
        "Many-to-One",
        "Some orders may have NULL customer",
    ]
    for col, val in enumerate(join_example, 1):
        cell = ws_joins.cell(row=5, column=col, value=val)
        cell.font = Font(italic=True, color="808080")
        style_data_cell(cell)

    for row in range(6, 15):
        for col in range(1, len(join_headers) + 1):
            style_data_cell(ws_joins.cell(row=row, column=col))

    set_column_widths(ws_joins, [8, 25, 25, 12, 25, 25, 20, 40])

    # ---------- VALIDATION SHEET ----------
    ws_validation = wb.create_sheet("Validation")

    ws_validation["A1"] = "VALIDATION PLAN"
    ws_validation["A1"].font = Font(bold=True, size=14)
    ws_validation["A2"] = "→ Write these BEFORE you build. How will you prove it works?"
    ws_validation["A2"].font = Font(italic=True, color="666666")

    ws_validation["A4"] = "ROW COUNT CHECKS"
    ws_validation["A4"].font = Font(bold=True)

    count_headers = ["Check", "Old Query/Source", "New Query/Target", "Expected", "Actual", "Pass?"]
    for col, header in enumerate(count_headers, 1):
        ws_validation.cell(row=5, column=col, value=header)
    style_header_row(ws_validation, 5, len(count_headers))

    for row in range(6, 10):
        for col in range(1, len(count_headers) + 1):
            style_data_cell(ws_validation.cell(row=row, column=col))

    ws_validation["A12"] = "SAMPLE RECORD CHECKS"
    ws_validation["A12"].font = Font(bold=True)
    ws_validation["A13"] = "→ Pick 3-5 specific records. Compare old vs new."
    ws_validation["A13"].font = Font(italic=True, color="666666")

    sample_headers = ["Record Key", "Field", "Old Value", "New Value", "Match?", "Notes"]
    for col, header in enumerate(sample_headers, 1):
        ws_validation.cell(row=14, column=col, value=header)
    style_header_row(ws_validation, 14, len(sample_headers))

    for row in range(15, 25):
        for col in range(1, len(sample_headers) + 1):
            style_data_cell(ws_validation.cell(row=row, column=col))

    ws_validation["A27"] = "BUSINESS RULE CHECKS"
    ws_validation["A27"].font = Font(bold=True)
    ws_validation["A28"] = (
        "→ What business rules must be true? (e.g., no negative amounts, all orders have customer)"
    )
    ws_validation["A28"].font = Font(italic=True, color="666666")

    rule_headers = ["Rule", "Validation Query", "Expected Result", "Actual", "Pass?"]
    for col, header in enumerate(rule_headers, 1):
        ws_validation.cell(row=29, column=col, value=header)
    style_header_row(ws_validation, 29, len(rule_headers))

    for row in range(30, 38):
        for col in range(1, len(rule_headers) + 1):
            style_data_cell(ws_validation.cell(row=row, column=col))

    set_column_widths(ws_validation, [25, 40, 40, 15, 10, 30])

    # ---------- OLD VS NEW SHEET ----------
    ws_compare = wb.create_sheet("Old vs New")

    ws_compare["A1"] = "OLD VS NEW COMPARISON"
    ws_compare["A1"].font = Font(bold=True, size=14)
    ws_compare["A2"] = "→ For migrations: side-by-side comparison of old and new approach"
    ws_compare["A2"].font = Font(italic=True, color="666666")

    ws_compare["A4"] = "OLD SQL / LOGIC"
    ws_compare["A4"].font = Font(bold=True)
    ws_compare["A5"] = "(Paste the old SQL here)"
    ws_compare.merge_cells("A5:F20")
    ws_compare["A5"].alignment = Alignment(vertical="top", wrap_text=True)

    ws_compare["A22"] = "NEW ODIBI APPROACH"
    ws_compare["A22"].font = Font(bold=True)
    ws_compare["A23"] = "(Describe or paste the Odibi config/approach here)"
    ws_compare.merge_cells("A23:F38")
    ws_compare["A23"].alignment = Alignment(vertical="top", wrap_text=True)

    ws_compare["A40"] = "KEY DIFFERENCES"
    ws_compare["A40"].font = Font(bold=True)

    diff_headers = ["Aspect", "Old Approach", "New Approach", "Why Changed"]
    for col, header in enumerate(diff_headers, 1):
        ws_compare.cell(row=41, column=col, value=header)
    style_header_row(ws_compare, 41, len(diff_headers))

    aspects = ["Join Logic", "Filtering", "Aggregation", "Error Handling", "Performance"]
    for i, aspect in enumerate(aspects):
        ws_compare.cell(row=42 + i, column=1, value=aspect)
        for col in range(1, 5):
            style_data_cell(ws_compare.cell(row=42 + i, column=col))

    set_column_widths(ws_compare, [20, 35, 35, 35])

    # ---------- CHECKLIST SHEET ----------
    ws_checklist = wb.create_sheet("Checklist")

    ws_checklist["A1"] = "COMPLETION CHECKLIST"
    ws_checklist["A1"].font = Font(bold=True, size=14)

    checklist_items = [
        (
            "PLANNING",
            [
                "Sources inventory complete",
                "Column mappings documented",
                "Join logic documented",
                "Validation queries written BEFORE building",
                "Old SQL/logic captured (if migration)",
            ],
        ),
        (
            "BUILD",
            [
                "Odibi config created",
                "Transformations implemented",
                "Unit tests written",
                "Local test passes",
            ],
        ),
        (
            "VALIDATE",
            [
                "Row counts match expected",
                "Sample records verified",
                "Business rules validated",
                "Edge cases tested (nulls, duplicates)",
                "Performance acceptable",
            ],
        ),
        (
            "FINALIZE",
            [
                "Documentation updated",
                "Master inventory updated",
                "Old process deprecated/removed",
                "Stakeholders notified",
            ],
        ),
    ]

    row = 3
    for section, items in checklist_items:
        ws_checklist.cell(row=row, column=1, value=section)
        ws_checklist.cell(row=row, column=1).font = Font(bold=True)
        row += 1
        for item in items:
            ws_checklist.cell(row=row, column=1, value="☐")
            ws_checklist.cell(row=row, column=2, value=item)
            row += 1
        row += 1

    set_column_widths(ws_checklist, [5, 50])

    wb.save(output_dir / "06_project_planning.xlsx")
    print("Created: 06_project_planning.xlsx")


def create_capture_log(output_dir):
    """Capture Log - Quick entries while working."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Capture Log"

    ws["A1"] = "CAPTURE LOG"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = "→ One-liners while working. Don't overthink, just capture. Process later."
    ws["A2"].font = Font(italic=True, color="666666")

    headers = ["Date", "Time", "What I Noticed", "Type", "Pipeline/Node", "Action Needed", "Status"]

    for col, header in enumerate(headers, 1):
        ws.cell(row=4, column=col, value=header)
    style_header_row(ws, 4, len(headers))

    # Example
    example = [
        "2025-01-06",
        "10:30",
        "SCD2 merge is slow on large tables",
        "Performance",
        "slv_customers",
        "Investigate partitioning",
        "Open",
    ]
    for col, val in enumerate(example, 1):
        cell = ws.cell(row=5, column=col, value=val)
        cell.font = Font(italic=True, color="808080")
        style_data_cell(cell)

    for row in range(6, 100):
        for col in range(1, len(headers) + 1):
            style_data_cell(ws.cell(row=row, column=col))

    set_column_widths(ws, [12, 8, 50, 15, 20, 30, 10])

    # Reference sheet
    ws2 = wb.create_sheet("Types")
    ws2["A1"] = "Type Options"
    ws2["A1"].font = Font(bold=True)
    types = [
        "Bug",
        "Idea",
        "Friction",
        "Win",
        "Question",
        "Performance",
        "Documentation",
        "Test Needed",
    ]
    for i, t in enumerate(types, 2):
        ws2.cell(row=i, column=1, value=t)

    ws2["C1"] = "Status Options"
    ws2["C1"].font = Font(bold=True)
    statuses = ["Open", "In Progress", "Done", "Won't Fix", "Captured for Content"]
    for i, s in enumerate(statuses, 2):
        ws2.cell(row=i, column=3, value=s)

    wb.save(output_dir / "07_capture_log.xlsx")
    print("Created: 07_capture_log.xlsx")


def create_content_bank(output_dir):
    """Content Bank - Ideas for LinkedIn/Medium campaign."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Content Ideas"

    ws["A1"] = "CONTENT IDEA BANK"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = "→ Capture ideas as they happen. This feeds your LinkedIn/Medium campaign."
    ws["A2"].font = Font(italic=True, color="666666")

    headers = [
        "Date Captured",
        "Idea/Hook",
        "Context/Story",
        "Type",
        "Phase/Week",
        "Status",
        "Published Date",
        "Link",
    ]

    for col, header in enumerate(headers, 1):
        ws.cell(row=4, column=col, value=header)
    style_header_row(ws, 4, len(headers))

    # Example
    example = [
        "2025-01-06",
        "Why I wrote my own SCD2 instead of using dbt",
        "Hit performance issues with dbt snapshots at scale, built custom solution",
        "LinkedIn Post",
        "Phase 3",
        "Draft",
        "",
        "",
    ]
    for col, val in enumerate(example, 1):
        cell = ws.cell(row=5, column=col, value=val)
        cell.font = Font(italic=True, color="808080")
        style_data_cell(cell)

    for row in range(6, 50):
        for col in range(1, len(headers) + 1):
            style_data_cell(ws.cell(row=row, column=col))

    set_column_widths(ws, [15, 40, 50, 15, 12, 12, 15, 30])

    # Reference sheet
    ws2 = wb.create_sheet("Reference")
    ws2["A1"] = "Content Types"
    ws2["A1"].font = Font(bold=True)
    types = ["LinkedIn Post", "Medium Article", "Twitter Thread", "Code Example", "Tutorial"]
    for i, t in enumerate(types, 2):
        ws2.cell(row=i, column=1, value=t)

    ws2["C1"] = "Status Options"
    ws2["C1"].font = Font(bold=True)
    statuses = ["Idea", "Outlined", "Draft", "Ready to Publish", "Published"]
    for i, s in enumerate(statuses, 2):
        ws2.cell(row=i, column=3, value=s)

    ws2["E1"] = "Campaign Phases"
    ws2["E1"].font = Font(bold=True)
    phases = [
        "Phase 1-2: Credibility",
        "Phase 3-10: Building in Public",
        "Phase 11-16: Pattern Deep Dives",
        "Phase 17-22: Anti-Patterns",
        "Phase 23-26: Advanced/Wrap-up",
    ]
    for i, p in enumerate(phases, 2):
        ws2.cell(row=i, column=5, value=p)

    wb.save(output_dir / "08_content_bank.xlsx")
    print("Created: 08_content_bank.xlsx")


def main():
    output_dir = Path(__file__).parent.parent / "templates"
    output_dir.mkdir(exist_ok=True)

    # Clear old templates
    for f in output_dir.glob("*.docx"):
        f.unlink()
    for f in output_dir.glob("*.xlsx"):
        f.unlink()

    print(f"Creating DE Planning System in: {output_dir}")
    print("=" * 60)
    print()

    # Word templates
    print("WORD TEMPLATES (Narrative Planning)")
    print("-" * 40)
    create_monthly_goals(output_dir)
    create_weekly_plan(output_dir)
    create_daily_plan(output_dir)
    print()

    # Excel templates
    print("EXCEL TEMPLATES (Structured Data)")
    print("-" * 40)
    create_session_backlog(output_dir)
    create_master_inventory(output_dir)
    create_project_planning(output_dir)
    create_capture_log(output_dir)
    create_content_bank(output_dir)
    print()

    print("=" * 60)
    print("Done! Complete DE Planning System created.")
    print()
    print("FILES CREATED:")
    print("  Word (narrative):")
    print("    - 01_monthly_goals.docx")
    print("    - 03_weekly_plan.docx")
    print("    - 04_daily_plan.docx")
    print()
    print("  Excel (structured):")
    print("    - 02_session_backlog.xlsx")
    print("    - 05_master_inventory.xlsx")
    print("    - 06_project_planning.xlsx (THE KEY ONE)")
    print("    - 07_capture_log.xlsx")
    print("    - 08_content_bank.xlsx")
    print()
    print("WORKFLOW:")
    print("  1. Monthly: Set goals, identify sessions needed")
    print("  2. Session Backlog: List all planned work")
    print("  3. Weekly: Allocate sessions to days")
    print("  4. Daily: Focus on today's sessions")
    print("  5. Project Planning: Deep planning for each pipeline")
    print("  6. Capture Log: Notes while working")
    print("  7. Content Bank: Ideas for marketing")


if __name__ == "__main__":
    main()
